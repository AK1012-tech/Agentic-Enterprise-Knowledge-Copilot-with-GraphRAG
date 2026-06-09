from __future__ import annotations

import csv
import io
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    text: str
    metadata: dict


class DocumentLoader:
    def parse(self, filename: str, payload: bytes) -> ParsedDocument:
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        if extension == "pdf":
            return self._parse_pdf(filename, payload)
        if extension == "docx":
            return self._parse_docx(filename, payload)
        if extension == "csv":
            return self._parse_csv(filename, payload)
        return ParsedDocument(text=payload.decode("utf-8", errors="ignore"), metadata={"source": filename})

    def _parse_pdf(self, filename: str, payload: bytes) -> ParsedDocument:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(payload))
            pages = [page.extract_text() or "" for page in reader.pages]
            return ParsedDocument(
                text="\n\n".join(pages),
                metadata={"source": filename, "pages": len(pages)},
            )
        except Exception:
            return ParsedDocument(text=payload.decode("utf-8", errors="ignore"), metadata={"source": filename})

    def _parse_docx(self, filename: str, payload: bytes) -> ParsedDocument:
        try:
            from docx import Document

            document = Document(io.BytesIO(payload))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            return ParsedDocument(text=text, metadata={"source": filename})
        except Exception:
            return ParsedDocument(text=payload.decode("utf-8", errors="ignore"), metadata={"source": filename})

    def _parse_csv(self, filename: str, payload: bytes) -> ParsedDocument:
        decoded = payload.decode("utf-8", errors="ignore")
        rows = csv.reader(io.StringIO(decoded))
        text = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
        return ParsedDocument(text=text, metadata={"source": filename})

